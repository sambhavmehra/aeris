"""
AERIS Word Transcript & Notes Generator Tools
Provides transcript/content extraction from YouTube and general web links,
premium-styled Word document (.docx) generation, and automated file opening.
"""
import os
import re
import sys
import logging
import platform
import subprocess
import urllib.parse
from pathlib import Path
from typing import Optional, Dict, List, Any

# Third-party libraries
import httpx
from bs4 import BeautifulSoup
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from youtube_transcript_api import YouTubeTranscriptApi

from config import settings
from ai_engine import ai_engine

logger = logging.getLogger("aeris.tools.word")


def get_backend_root() -> Path:
    """Helper to resolve the BACKEND directory root."""
    return Path(__file__).resolve().parent.parent


def get_workspace_dir() -> Path:
    """Get the active workspace directory path."""
    return Path(settings.WORKSPACE_DIR)


def extract_youtube_video_id(url: str) -> Optional[str]:
    """
    Extract the 11-character YouTube video ID from various URL formats.
    Supports watch URLs, share URLs, embeds, shorts, and mobile formats.
    """
    url = url.strip()
    
    # If the URL is already just an 11-character string that looks like a video ID
    if len(url) == 11 and re.match(r'^[A-Za-z0-9_-]{11}$', url):
        return url
        
    # Regex patterns to capture video ID
    patterns = [
        r"(?:v=|\/v\/|embed\/|shorts\/|youtu\.be\/|\/embed\/|\/v=|^v$|^video_id$)([^#\&\?]*)"
    ]
    for pattern in patterns:
        match = re.search(pattern, url)
        if match:
            video_id = match.group(1)
            if len(video_id) == 11:
                return video_id

    # Fallback using urllib parsing
    try:
        parsed_url = urllib.parse.urlparse(url)
        if parsed_url.hostname in ('www.youtube.com', 'youtube.com', 'm.youtube.com', 'youtube-nocookie.com'):
            query_params = urllib.parse.parse_qs(parsed_url.query)
            if 'v' in query_params:
                return query_params['v'][0]
        elif parsed_url.hostname == 'youtu.be':
            path = parsed_url.path.strip('/')
            if len(path) == 11:
                return path
    except Exception as e:
        logger.warning(f"Error parsing YouTube URL with urllib: {e}")
        
    return None


def get_youtube_transcript(video_id: str) -> str:
    """
    Retrieve transcript for a YouTube video using youtube-transcript-api.
    Prefers English, but falls back to other available languages.
    """
    # 1. Try classmethod style (standard library behavior for older/stable versions)
    if hasattr(YouTubeTranscriptApi, 'get_transcript'):
        try:
            transcript_list = YouTubeTranscriptApi.get_transcript(video_id)
            return " ".join([t['text'] for t in transcript_list])
        except Exception as e:
            logger.warning(f"First attempt to fetch YouTube transcript (classmethod) failed: {e}. Trying fallback.")
            if hasattr(YouTubeTranscriptApi, 'list_transcripts'):
                try:
                    transcript_list = YouTubeTranscriptApi.list_transcripts(video_id)
                    for tr in transcript_list:
                        return " ".join([t['text'] for t in tr.fetch()])
                except Exception as list_err:
                    raise RuntimeError(
                        f"YouTube transcript extraction failed: {e}. (Make sure the video exists and has public captions/subtitles enabled)."
                    ) from list_err

    # 2. Try instance style (required for newer/custom version 1.2.4)
    try:
        import requests
        session = requests.Session()
        session.headers.update({
            "User-Agent": (
                "Mozilla/5.0 (Windows NT 10.0; Win64; x64) "
                "AppleWebKit/537.36 (KHTML, like Gecko) "
                "Chrome/124.0.0.0 Safari/537.36"
            )
        })
        api = YouTubeTranscriptApi(http_client=session)
        try:
            transcript_list = api.fetch(video_id)
            return " ".join([t['text'] for t in transcript_list])
        except Exception as e:
            logger.warning(f"First attempt to fetch YouTube transcript (instance) failed: {e}. Trying fallback.")
            try:
                transcript_list = api.list(video_id)
                for tr in transcript_list:
                    return " ".join([t['text'] for t in tr.fetch()])
            except Exception as list_err:
                raise RuntimeError(
                    f"YouTube transcript extraction failed: {e}. (Make sure the video exists and has public captions/subtitles enabled)."
                ) from list_err
    except Exception as api_err:
        raise RuntimeError(
            f"YouTube transcript extraction failed: {api_err}. (Make sure the video exists and has public captions/subtitles enabled)."
        ) from api_err


def scrape_web_link(url: str) -> str:
    """
    Scrape text content from a general web page.
    Cleans up headers, footers, scripts, and navigation sections.
    """
    headers = {
        "User-Agent": (
            "Mozilla/5.0 (Windows NT 10.0; Win64; x64) "
            "AppleWebKit/537.36 (KHTML, like Gecko) "
            "Chrome/124.0.0.0 Safari/537.36"
        ),
    }
    try:
        with httpx.Client(timeout=15.0, follow_redirects=True, headers=headers) as client:
            resp = client.get(url)
            resp.raise_for_status()
            html = resp.text

        soup = BeautifulSoup(html, "html.parser")
        
        # Decompose non-content tags
        for tag in soup(["script", "style", "nav", "footer", "header", "aside", "form", "noscript", "iframe", "svg"]):
            tag.decompose()

        # Extract textual content
        raw_text = soup.get_text(separator="\n", strip=True)
        lines = [l.strip() for l in raw_text.splitlines() if l.strip()]
        clean_text = "\n".join(lines)
        return clean_text
    except Exception as e:
        raise RuntimeError(f"Webpage scraping failed: {e}") from e


def add_markdown_paragraph(doc: Document, md_line: str) -> None:
    """
    Parse a single markdown line and add it as a styled paragraph/heading/list to the docx.
    """
    md_line = md_line.strip()
    if not md_line:
        return

    # 1. Handle Headings
    if md_line.startswith("# "):
        p = doc.add_paragraph()
        run = p.add_run(md_line[2:])
        run.font.name = "Segoe UI"
        run.font.size = Pt(20)
        run.font.bold = True
        run.font.color.rgb = RGBColor(30, 63, 102) # Slate Blue
        p.paragraph_format.space_before = Pt(14)
        p.paragraph_format.space_after = Pt(6)
        p.paragraph_format.keep_with_next = True
        return
    elif md_line.startswith("## "):
        p = doc.add_paragraph()
        run = p.add_run(md_line[3:])
        run.font.name = "Segoe UI"
        run.font.size = Pt(15)
        run.font.bold = True
        run.font.color.rgb = RGBColor(30, 63, 102) # Slate Blue
        p.paragraph_format.space_before = Pt(12)
        p.paragraph_format.space_after = Pt(4)
        p.paragraph_format.keep_with_next = True
        return
    elif md_line.startswith("### "):
        p = doc.add_paragraph()
        run = p.add_run(md_line[4:])
        run.font.name = "Segoe UI"
        run.font.size = Pt(13)
        run.font.bold = True
        run.font.color.rgb = RGBColor(56, 128, 135) # Accent Teal
        p.paragraph_format.space_before = Pt(10)
        p.paragraph_format.space_after = Pt(4)
        p.paragraph_format.keep_with_next = True
        return

    # 2. Handle Lists and Blockquotes
    is_bullet = False
    is_numbered = False
    is_blockquote = False
    list_num_prefix = ""
    
    if md_line.startswith("- ") or md_line.startswith("* "):
        is_bullet = True
        content = md_line[2:]
    elif md_line.startswith("> "):
        is_blockquote = True
        content = md_line[2:]
    elif re.match(r'^\d+\.\s+', md_line):
        is_numbered = True
        parts = md_line.split(".", 1)
        list_num_prefix = parts[0] + "."
        content = parts[1].strip()
    else:
        content = md_line

    # Create paragraph with appropriate style
    if is_bullet:
        p = doc.add_paragraph(style='List Bullet')
    elif is_blockquote:
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Inches(0.5)
    else:
        p = doc.add_paragraph()

    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.line_spacing = 1.15

    # If numbered, add prefix first
    if is_numbered:
        prefix_run = p.add_run(list_num_prefix + " ")
        prefix_run.font.name = "Calibri"
        prefix_run.font.bold = True
        prefix_run.font.color.rgb = RGBColor(30, 63, 102)

    # 3. Parse Inline Formatting: ***bold-italic***, **bold**, *italic*, _italic_, `code`
    # Pattern groups formatting markers
    pattern = re.compile(r'(\*\*\*.*?\*\*\*|\*\*.*?\*\*|\*.*?\*|`.*?`|_+.*?_+)')
    parts = pattern.split(content)

    for part in parts:
        if not part:
            continue
        
        # Check matching style and extract raw text
        if part.startswith("***") and part.endswith("***"):
            run = p.add_run(part[3:-3])
            run.font.bold = True
            run.font.italic = True
        elif part.startswith("**") and part.endswith("**"):
            run = p.add_run(part[2:-2])
            run.font.bold = True
        elif part.startswith("*") and part.endswith("*"):
            run = p.add_run(part[1:-1])
            run.font.italic = True
        elif part.startswith("_") and part.endswith("_"):
            run = p.add_run(part[1:-1])
            run.font.italic = True
        elif part.startswith("`") and part.endswith("`"):
            run = p.add_run(part[1:-1])
            run.font.name = "Consolas"
            run.font.size = Pt(9.5)
            run.font.color.rgb = RGBColor(180, 40, 40) # Muted red for code
        else:
            run = p.add_run(part)
            if is_blockquote:
                run.font.italic = True
                run.font.color.rgb = RGBColor(100, 100, 100)

        # Set default font name
        if not part.startswith("`"):
            run.font.name = "Calibri"
            run.font.size = Pt(11)


async def extract_transcript_to_word(
    url: Optional[str] = None,
    text: Optional[str] = None,
    output_path: Optional[str] = None,
    notes_style: str = "detailed"
) -> str:
    """
    Main tool entry point:
    1. Extracts transcript/text from a YouTube link, general link, or takes raw text.
    2. Uses AERIS AI Engine to generate detailed structured notes/summaries.
    3. Generates a beautifully formatted Microsoft Word (.docx) document.
    4. Automatically saves and opens the file.
    """
    if not url and not text:
        raise ValueError("Either 'url' or 'text' parameter must be provided.")

    workspace = get_workspace_dir()
    source_title = "Document"
    extracted_content = ""
    yt_error_msg = None

    # Step 1: Content Extraction
    if url:
        url_lower = url.lower().strip()
        is_youtube = "youtube.com" in url_lower or "youtu.be" in url_lower
        
        if is_youtube:
            video_id = extract_youtube_video_id(url)
            if not video_id:
                raise ValueError(f"Could not extract a valid 11-character YouTube Video ID from: {url}")
            
            logger.info(f"Extracting YouTube transcript for video ID: {video_id}")
            try:
                extracted_content = get_youtube_transcript(video_id)
            except Exception as yt_err:
                logger.warning(f"YouTube transcript fetch failed: {yt_err}. Trying metadata extraction via yt-dlp as fallback.")
                yt_error_msg = str(yt_err)
                try:
                    import yt_dlp
                    ydl = yt_dlp.YoutubeDL({'skip_download': True, 'quiet': True})
                    info = ydl.extract_info(url, download=False)
                    extracted_content = (
                        f"[SYSTEM NOTICE: YouTube transcript download was blocked/unavailable (Error: {yt_err}). "
                        f"Fell back to extracting YouTube video metadata.]\n\n"
                        f"Video Title: {info.get('title')}\n"
                        f"Channel: {info.get('uploader')}\n"
                        f"Description:\n{info.get('description')}\n"
                    )
                except Exception as ytdl_err:
                    logger.warning(f"yt-dlp metadata fetch failed: {ytdl_err}. Trying general web scraping fallback.")
                    try:
                        page_html = scrape_web_link(url)
                        extracted_content = (
                            f"[SYSTEM NOTICE: YouTube transcript download was blocked/unavailable (Error: {yt_err}). "
                            f"Fell back to scraping YouTube video watch page HTML.]\n\n"
                            f"{page_html}"
                        )
                    except Exception as scrape_err:
                        raise RuntimeError(
                            f"YouTube transcript extraction failed: {yt_err}. "
                            f"Scraping watch page as fallback also failed: {scrape_err}"
                        ) from scrape_err
            source_title = f"YouTube Video ({video_id})"
        else:
            logger.info(f"Scraping general web link: {url}")
            extracted_content = scrape_web_link(url)
            # Try to grab domain name as title
            try:
                parsed = urllib.parse.urlparse(url)
                source_title = f"Web Article ({parsed.netloc})"
            except Exception:
                source_title = "Web Article"
    else:
        logger.info("Using raw text provided directly by user")
        extracted_content = text
        source_title = "User Provided Text"

    if not extracted_content or not extracted_content.strip():
        raise RuntimeError("Failed to extract or retrieve any content to summarize.")

    # Limit size of prompt content for reliability
    content_snippet = extracted_content[:30000]

    # Step 2: Query AI Engine for Summarization/Notes
    logger.info("Querying AI Engine to generate structured notes")
    system_prompt = (
        "You are an expert professional notes taker and summary writer.\n"
        "Your task is to analyze the provided text or transcript and generate highly detailed, "
        "comprehensive, and structured study or meeting notes.\n"
        "Structure your response strictly in clear sections using Markdown:\n\n"
        "# Executive Summary\n"
        "(A solid paragraph summarizing the core theme, context, and purpose)\n\n"
        "# Key Takeaways & Insights\n"
        "(Bullet points detailing the most critical lessons, facts, or results)\n\n"
        "# Comprehensive Notes\n"
        "(Detailed sub-sections mapping different parts/themes, summarizing details in-depth)\n\n"
        "# Action Items & Next Steps\n"
        "(Numbered list of actions or future directions. If none apply, list recommendations instead)\n\n"
        "# Conclusion\n"
        "(Brief closing/wrap-up statement)\n\n"
        "Formatting Rules:\n"
        "- Use proper markdown headings (# for major titles, ## for sections, ### for sub-headers).\n"
        "- Bold important keywords or phrases to make the notes easy to scan.\n"
        "- Keep the tone professional, clear, and highly informative."
    )
    
    user_prompt = (
        f"Generate notes based on the following content extracted from {source_title}:\n\n"
        f"--- START CONTENT ---\n{content_snippet}\n--- END CONTENT ---"
    )

    try:
        notes_md = await ai_engine.chat([
            {"role": "system", "content": system_prompt},
            {"role": "user", "content": user_prompt}
        ], max_tokens=3000)
    except Exception as llm_err:
        logger.error(f"Failed to generate notes using AI Engine: {llm_err}")
        # Local fallback notes formatting if LLM fails
        notes_md = (
            f"# Notes: {source_title}\n\n"
            f"## Summary\n"
            f"Note generation failed, but content was successfully extracted. "
            f"Here is a snippet of the extracted raw content:\n\n"
            f"> {content_snippet[:1000]}..."
        )

    # Step 3: Parse and generate .docx file
    logger.info("Initializing docx generation")
    doc = Document()
    
    # Configure page margins
    for s in doc.sections:
        s.top_margin = Inches(1)
        s.bottom_margin = Inches(1)
        s.left_margin = Inches(1)
        s.right_margin = Inches(1)

    # Document Header Style Title
    title_p = doc.add_paragraph()
    title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_run = title_p.add_run("AERIS KNOWLEDGE NOTEBOOK")
    title_run.font.name = "Segoe UI"
    title_run.font.size = Pt(24)
    title_run.font.bold = True
    title_run.font.color.rgb = RGBColor(30, 63, 102) # Slate Blue
    title_p.paragraph_format.space_after = Pt(4)

    # Subtitle with source metadata
    sub_p = doc.add_paragraph()
    sub_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub_run = sub_p.add_run(f"Structured Notes generated from {source_title}")
    sub_run.font.name = "Calibri"
    sub_run.font.size = Pt(11)
    sub_run.font.italic = True
    sub_run.font.color.rgb = RGBColor(120, 120, 120)
    sub_p.paragraph_format.space_after = Pt(24)

    # Parse notes markdown and add elements
    if yt_error_msg:
        warn_p = doc.add_paragraph()
        warn_run = warn_p.add_run(
            "[WARNING] NOTICE: YouTube is currently rate-limiting or blocking automatic transcript downloads on your IP address. "
            "AERIS has fallen back to scraping the video's title, description, and page metadata to generate these notes. "
            "If you need notes on the full spoken content, please open the video on YouTube, open the Transcript panel, "
            "copy the text, and run this tool using the 'text' parameter."
        )
        warn_run.font.name = "Calibri"
        warn_run.font.size = Pt(10.5)
        warn_run.font.italic = True
        warn_run.font.color.rgb = RGBColor(180, 50, 50)
        warn_p.paragraph_format.space_after = Pt(12)

    for line in notes_md.splitlines():
        add_markdown_paragraph(doc, line)

    # Determine save path
    if output_path:
        resolved_path = Path(output_path)
        if not resolved_path.is_absolute():
            resolved_path = workspace / output_path
    else:
        # Generate safe filename from source
        safe_title = re.sub(r"[^A-Za-z0-9._-]+", "_", source_title.lower())[:30]
        resolved_path = workspace / f"notes_{safe_title}.docx"

    resolved_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(resolved_path))
    logger.info(f"Saved styled notes Word document at: {resolved_path}")

    # Track file creation in AERIS
    try:
        from utils.file_tracker import record_file_creation
        record_file_creation(str(resolved_path), f"Generated styled Word document notes from {source_title}")
    except Exception:
        pass

    # Step 4: Open File Automatically
    opened_status = "automatically opened"
    try:
        if platform.system() == "Windows":
            os.startfile(str(resolved_path))
        elif platform.system() == "Darwin":
            subprocess.Popen(["open", str(resolved_path)])
        else:
            subprocess.Popen(["xdg-open", str(resolved_path)])
    except Exception as open_err:
        logger.warning(f"Failed to automatically open the document: {open_err}")
        opened_status = "saved (but could not be opened automatically due to system settings)"

    warning_notice = ""
    if yt_error_msg:
        warning_notice = (
            "\n\n[WARNING] **Note regarding YouTube Transcript:** YouTube is currently rate-limiting transcript downloads on your IP address (429/IpBlocked). "
            "I have gracefully fallen back to extracting the video's title, description, and metadata from the watch page to generate these notes. "
            "If you need notes on the full spoken content, please open the video on YouTube, open the Transcript panel, copy the text, "
            "and paste it to me directly using the 'text' option!"
        )

    return (
        f"Sir, I have successfully extracted the content from '{source_title}', "
        f"generated detailed structured notes using AI, and saved the document in Word format at:\n"
        f"`{resolved_path}`.{warning_notice}\n\n"
        f"The file has been {opened_status} for you."
    )
