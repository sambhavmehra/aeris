"""
AERIS — Agentic File & Bash Tool System
Real implementations for read_file, write_file, edit_file, delete_file,
glob_search, and bash execution with sandbox enforcement.
"""
from __future__ import annotations

import glob
import json
import os
import shutil
import logging
from dataclasses import dataclass
from datetime import datetime
from pathlib import Path
from typing import Optional

from services.permission_enforcer import PermissionEnforcer, SandboxExecutor
from services.path_utils import resolve_path, get_shared_directories, list_shared_dirs_status

logger = logging.getLogger("FileTools")


@dataclass(frozen=True)
class FileToolResult:
    success: bool
    output: str
    tool: str


class FileToolSystem:
    """Real file system operations with permission enforcement."""

    def __init__(self, workspace: str | None = None) -> None:
        self.workspace = Path(workspace or os.getcwd()).resolve()
        self.enforcer = PermissionEnforcer(workspace_root=self.workspace)
        self.sandbox = SandboxExecutor(enforcer=self.enforcer)

    # ── read_file ────────────────────────────────────────────────────
    def read_file(self, filepath: str) -> FileToolResult:
        filepath = os.path.expanduser(filepath)
        target = (self.workspace / filepath).resolve() if not Path(filepath).is_absolute() else Path(filepath).resolve()
        check = self.enforcer.check_tool("read_file")
        if not check.allowed:
            return FileToolResult(False, check.reason, "read_file")
        try:
            if not target.exists():
                return FileToolResult(False, f"File not found: {target}", "read_file")
            if target.stat().st_size > 10_000_000:  # 10MB limit
                return FileToolResult(False, "File too large (>10MB)", "read_file")
            # Binary detection
            sample = target.read_bytes()[:8192]
            if b"\x00" in sample:
                extracted_text = self._extract_text_from_binary(target)
                if extracted_text is not None:
                    return FileToolResult(True, extracted_text, "read_file")
                return FileToolResult(False, "Binary file detected — refusing to read", "read_file")
            content = target.read_text(encoding="utf-8", errors="replace")
            return FileToolResult(True, content, "read_file")
        except Exception as exc:
            return FileToolResult(False, str(exc), "read_file")

    # ── write_file ───────────────────────────────────────────────────
    def write_file(self, filepath: str, content: str) -> FileToolResult:
        filepath = os.path.expanduser(filepath)
        target = (self.workspace / filepath).resolve() if not Path(filepath).is_absolute() else Path(filepath).resolve()
        check = self.enforcer.check_file_write(str(target))
        if not check.allowed:
            return FileToolResult(False, check.reason, "write_file")
        try:
            target.parent.mkdir(parents=True, exist_ok=True)
            target.write_text(content, encoding="utf-8")
            return FileToolResult(True, f"Written {len(content)} bytes to {target}", "write_file")
        except Exception as exc:
            return FileToolResult(False, str(exc), "write_file")

    # ── edit_file ────────────────────────────────────────────────────
    def edit_file(self, filepath: str, old_text: str, new_text: str) -> FileToolResult:
        filepath = os.path.expanduser(filepath)
        target = (self.workspace / filepath).resolve() if not Path(filepath).is_absolute() else Path(filepath).resolve()
        check = self.enforcer.check_file_write(str(target))
        if not check.allowed:
            return FileToolResult(False, check.reason, "edit_file")
        ext = target.suffix.lower()
        try:
            if not target.exists():
                return FileToolResult(False, f"File not found: {target}", "edit_file")

            if ext in (".docx", ".doc"):
                try:
                    from docx import Document
                    doc = Document(str(target))
                    found = False
                    for para in doc.paragraphs:
                        if old_text in para.text:
                            para.text = para.text.replace(old_text, new_text, 1)
                            found = True
                            break
                    if not found:
                        return FileToolResult(False, "Target text not found in DOCX file", "edit_file")
                    doc.save(str(target))
                    return FileToolResult(True, f"Edited DOCX file {target}", "edit_file")
                except ImportError:
                    return FileToolResult(False, "Missing dependency: python-docx. Cannot edit DOCX.", "edit_file")

            content = target.read_text(encoding="utf-8")
            if old_text not in content:
                return FileToolResult(False, "Target text not found in file", "edit_file")
            updated = content.replace(old_text, new_text, 1)
            target.write_text(updated, encoding="utf-8")
            return FileToolResult(True, f"Edited {target}", "edit_file")
        except Exception as exc:
            return FileToolResult(False, str(exc), "edit_file")

    # ── delete_file ──────────────────────────────────────────────────
    def delete_file(self, filepath: str) -> FileToolResult:
        filepath = os.path.expanduser(filepath)
        target = (self.workspace / filepath).resolve() if not Path(filepath).is_absolute() else Path(filepath).resolve()
        check = self.enforcer.check_file_write(str(target))
        if not check.allowed:
            return FileToolResult(False, check.reason, "delete_file")
        try:
            if target.is_file():
                target.unlink()
                return FileToolResult(True, f"Deleted {target}", "delete_file")
            elif target.is_dir():
                shutil.rmtree(target)
                return FileToolResult(True, f"Deleted directory {target}", "delete_file")
            else:
                return FileToolResult(False, f"Not found: {target}", "delete_file")
        except Exception as exc:
            return FileToolResult(False, str(exc), "delete_file")

    # ── glob_search ──────────────────────────────────────────────────
    def glob_search(self, pattern: str, max_results: int = 100) -> FileToolResult:
        try:
            matches = list(self.workspace.glob(pattern))[:max_results]
            relative = [str(m.relative_to(self.workspace)) for m in matches]
            return FileToolResult(
                True,
                "\n".join(relative) if relative else "No matches found",
                "glob_search",
            )
        except Exception as exc:
            return FileToolResult(False, str(exc), "glob_search")

    # ── grep_search ──────────────────────────────────────────────────
    def grep_search(self, query: str, path: str = ".", max_results: int = 50) -> FileToolResult:
        target = (self.workspace / path).resolve()
        try:
            results: list[str] = []
            search_paths = target.rglob("*") if target.is_dir() else [target]
            for fpath in search_paths:
                if not fpath.is_file() or fpath.stat().st_size > 5_000_000:
                    continue
                try:
                    for i, line in enumerate(fpath.read_text(errors="ignore").splitlines(), 1):
                        if query.lower() in line.lower():
                            rel = fpath.relative_to(self.workspace)
                            results.append(f"{rel}:{i}: {line.strip()}")
                            if len(results) >= max_results:
                                break
                except (UnicodeDecodeError, PermissionError):
                    continue
                if len(results) >= max_results:
                    break
            return FileToolResult(
                True,
                "\n".join(results) if results else "No matches found",
                "grep_search",
            )
        except Exception as exc:
            return FileToolResult(False, str(exc), "grep_search")

    # ── bash ─────────────────────────────────────────────────────────
    def bash(self, command: str) -> FileToolResult:
        import re
        import tempfile
        import os

        command = command.strip()

        # ── FIX 1: Expand tilde in commands ──────────────────────────
        # Replace ~/ or ~/path with the actual home directory
        # Quotes unquoted paths automatically to handle user directories with spaces
        home_path = os.path.expanduser("~").replace("\\", "/")

        # Map ~/Desktop and ~/Documents to OneDrive if they exist
        onedrive_desktop = os.path.join(os.path.expanduser("~"), "OneDrive", "Desktop")
        onedrive_docs = os.path.join(os.path.expanduser("~"), "OneDrive", "Documents")
        if os.path.exists(onedrive_desktop):
            command = re.sub(r'~/Desktop(?=/|$)', "~/OneDrive/Desktop", command)
        if os.path.exists(onedrive_docs):
            command = re.sub(r'~/Documents(?=/|$)', "~/OneDrive/Documents", command)

        command = re.sub(
            r'(?<!["\'\S])(~/[^\s"\']*)',
            lambda m: f'"{home_path}{m.group(1)[1:]}"',
            command
        )
        command = re.sub(
            r'(?<=["\'])~/',
            home_path + "/",
            command
        )

        # ── FIX 2: python -c '...' or python -c "..." → temp file ───
        # Windows PowerShell mangles both single and double quotes in
        # inline python commands.  We intercept them, write to a temp
        # file and execute that instead.
        tmp_path = None
        match = re.match(
            r'^python\s+-c\s+["\'](.+)["\']$', command, re.DOTALL
        )
        if match:
            code = match.group(1)
            fd, tmp_path = tempfile.mkstemp(suffix=".py")
            with os.fdopen(fd, "w", encoding="utf-8") as f:
                f.write(code)
            command = f'python "{tmp_path}"'

        # ── FIX 3: Auto-quote unquoted paths with spaces ─────────────
        # Catches:  python  D:\Some Folder\script.py  [args...]
        #           python  Documents\my script.py
        # and wraps just the path portion in quotes so PowerShell
        # doesn't split it on the space.
        if not tmp_path:
            m = re.match(
                r'^(python\d?(?:\.exe)?)\s+'       # python executable
                r'(?!-)'                             # not a flag like -c, -m
                r'([^"\']\S*(?:\s+\S+)*?\.py\b)'   # unquoted .py path
                r'((?:\s+.*)?)$',                   # optional trailing args
                command,
                re.IGNORECASE,
            )
            if m:
                exe, script_path, rest = m.group(1), m.group(2), m.group(3)
                # Only quote if there's actually a space in the path
                if " " in script_path and not script_path.startswith('"'):
                    command = f'{exe} "{script_path}"{rest}'

        # ── Execute via Sandbox (security enforced) ──────────────────
        result = self.sandbox.execute_bash(command, cwd=str(self.workspace))

        if tmp_path:
            try:
                os.remove(tmp_path)
            except Exception:
                pass

        if result["blocked"]:
            return FileToolResult(False, result["reason"], "bash")
        output = result["stdout"]
        if result["stderr"]:
            output += f"\n[stderr] {result['stderr']}"
        return FileToolResult(result["success"], output, "bash")

    # ── read_system_file (any absolute path) ─────────────────────────
    def read_system_file(self, filepath: str) -> FileToolResult:
        """Read any file on the system by absolute path. Not restricted to workspace."""
        filepath = os.path.expanduser(filepath)
        target = Path(filepath).resolve()
        try:
            if not target.exists():
                return FileToolResult(False, f"File not found: {target}", "read_system_file")
            if not target.is_file():
                return FileToolResult(False, f"Not a file: {target}", "read_system_file")
            if target.stat().st_size > 10_000_000:  # 10MB limit
                return FileToolResult(False, f"File too large (>10MB): {target.stat().st_size} bytes", "read_system_file")
            # Binary file detection
            sample = target.read_bytes()[:8192]
            if b"\x00" in sample:
                extracted_text = self._extract_text_from_binary(target)
                if extracted_text is not None:
                    return FileToolResult(True, extracted_text, "read_system_file")
                return FileToolResult(
                    True,
                    f"Binary file detected ({target.stat().st_size} bytes). Path: {target}",
                    "read_system_file",
                )
            content = target.read_text(encoding="utf-8", errors="replace")
            return FileToolResult(True, content, "read_system_file")
        except PermissionError:
            return FileToolResult(False, f"Permission denied: {target}", "read_system_file")
        except Exception as exc:
            return FileToolResult(False, str(exc), "read_system_file")

    # ── list_system_dir (any absolute path) ──────────────────────────
    def list_system_dir(self, dirpath: str) -> FileToolResult:
        """List contents of any directory on the system by absolute path."""
        dirpath = os.path.expanduser(dirpath)
        target = Path(dirpath).resolve()
        try:
            if not target.exists():
                return FileToolResult(False, f"Directory not found: {target}", "list_system_dir")
            if not target.is_dir():
                return FileToolResult(False, f"Not a directory: {target}", "list_system_dir")

            entries = []
            for entry in sorted(target.iterdir()):
                try:
                    if entry.is_dir():
                        count = sum(1 for _ in entry.iterdir()) if entry.is_dir() else 0
                        entries.append(f"[DIR]  {entry.name}/ ({count} items)")
                    else:
                        size = entry.stat().st_size
                        if size < 1024:
                            size_str = f"{size} B"
                        elif size < 1024 * 1024:
                            size_str = f"{size / 1024:.1f} KB"
                        else:
                            size_str = f"{size / (1024 * 1024):.1f} MB"
                        entries.append(f"[FILE] {entry.name} ({size_str})")
                except PermissionError:
                    entries.append(f"[????] {entry.name} (access denied)")
                except Exception:
                    entries.append(f"[????] {entry.name}")

            if not entries:
                return FileToolResult(True, f"Directory is empty: {target}", "list_system_dir")

            header = f"Contents of {target} ({len(entries)} items):\n"
            return FileToolResult(True, header + "\n".join(entries), "list_system_dir")
        except PermissionError:
            return FileToolResult(False, f"Permission denied: {target}", "list_system_dir")
        except Exception as exc:
            return FileToolResult(False, str(exc), "list_system_dir")

    # ── get_file_info (metadata about any file) ──────────────────────
    def get_file_info(self, filepath: str) -> FileToolResult:
        """Get detailed metadata about any file or directory on the system."""
        filepath = os.path.expanduser(filepath)
        target = Path(filepath).resolve()
        try:
            if not target.exists():
                return FileToolResult(False, f"Path not found: {target}", "get_file_info")

            stat = target.stat()
            info = {
                "path": str(target),
                "name": target.name,
                "extension": target.suffix,
                "is_file": target.is_file(),
                "is_dir": target.is_dir(),
                "size_bytes": stat.st_size,
                "size_human": self._human_size(stat.st_size),
                "created": datetime.fromtimestamp(stat.st_ctime).isoformat(),
                "modified": datetime.fromtimestamp(stat.st_mtime).isoformat(),
                "parent": str(target.parent),
            }

            if target.is_dir():
                try:
                    children = list(target.iterdir())
                    info["child_count"] = len(children)
                    info["files"] = sum(1 for c in children if c.is_file())
                    info["subdirs"] = sum(1 for c in children if c.is_dir())
                except PermissionError:
                    info["child_count"] = "access denied"

            import json
            return FileToolResult(True, json.dumps(info, indent=2), "get_file_info")
        except PermissionError:
            return FileToolResult(False, f"Permission denied: {target}", "get_file_info")
        except Exception as exc:
            return FileToolResult(False, str(exc), "get_file_info")

    # ── find_system_file (search by name anywhere) ───────────────────
    # Folders to skip during search (noisy / system / sync folders)
    _SKIP_DIRS = frozenset({
        'AppData', 'Windows', 'Program Files', 'Program Files (x86)',
        'node_modules', '__pycache__', '.git', 'CrossDevice', 'ansel',
        'Saved Games', '.cache', '.ollama', '.docker', '.vscode',
        '.conda', '.npm', '.nuget', 'ProgramData',
    })

    def find_system_file(self, filename: str, search_dir: str = "~") -> FileToolResult:
        """Search for a file by name. Prioritizes Downloads, Documents, Desktop."""
        # Use path_utils for consistent resolution
        try:
            start_dir = resolve_path(search_dir) if search_dir else Path.home()
        except Exception as e:
            logger.error(f"Failed to resolve search directory '{search_dir}': {e}")
            start_dir = Path.home()
        
        try:
            if not start_dir.exists() or not start_dir.is_dir():
                return FileToolResult(False, f"Search directory invalid: {start_dir}", "find_system_file")

            logger.info(f"🔍 Searching for file '{filename}' starting from: {start_dir}")

            # Build search terms from the query
            q_lower = filename.lower().strip()
            # Extract the file extension if present
            q_ext = ""
            for ext in (".pdf", ".docx", ".doc", ".txt", ".xlsx", ".csv", ".pptx", ".png", ".jpg", ".mp4"):
                if q_lower.endswith(ext):
                    q_ext = ext
                    break
            q_no_ext = q_lower.removesuffix(q_ext) if q_ext else q_lower
            terms = [t for t in q_no_ext.split() if len(t) >= 2]

            # Scoring function: higher = better match
            def _score(filepath: str) -> int:
                fname = os.path.basename(filepath).lower()
                s = 0
                # Exact filename match (with or without spaces)
                if fname == q_lower or fname == q_lower.replace(" ", ""):
                    s += 1000
                # Extension match
                if q_ext and fname.endswith(q_ext):
                    s += 100
                # All terms present
                if terms and all(t in fname for t in terms):
                    s += 50
                # Bonus for being in a priority folder
                fp_lower = filepath.lower()
                for pf in ("downloads", "documents", "desktop", "onedrive"):
                    if pf in fp_lower:
                        s += 30
                        break
                # Penalty for deep nesting
                depth = filepath.count(os.sep)
                s -= depth
                return s

            # Priority folders to search first
            priority_names = ["Downloads", "Documents", "Desktop", "OneDrive"]
            priority_dirs = [start_dir / pn for pn in priority_names if (start_dir / pn).is_dir()]
            
            # Add workspace to priority directories if it exists
            if self.workspace and self.workspace.is_dir() and self.workspace not in priority_dirs:
                priority_dirs.insert(0, self.workspace)
            
            # Log which priority directories are being searched
            for pdir in priority_dirs:
                logger.info(f"  📁 Searching in: {pdir}")
            
            other_dirs = [
                d for d in start_dir.iterdir()
                if d.is_dir() and d.name not in priority_names
                and d.name not in self._SKIP_DIRS and not d.name.startswith('.')
            ]
            search_order = priority_dirs + other_dirs

            matches: list[str] = []

            for search_root in search_order:
                for root, dirs, files in os.walk(str(search_root)):
                    dirs[:] = [d for d in dirs if d not in self._SKIP_DIRS and not d.startswith('.')]
                    for file in files:
                        f_lower = file.lower()
                        f_clean = f_lower.replace(" ", "")
                        q_clean = q_lower.replace(" ", "")

                        # If query has an extension (e.g. .pdf, .doc), the matched file MUST end with that extension
                        if q_ext and not f_lower.endswith(q_ext):
                            continue

                        is_match = False
                        # 1. Exact / substring match (ignoring spaces)
                        if q_clean in f_clean or f_clean in q_clean:
                            is_match = True
                        # 2. All search terms present in filename
                        elif terms and all(t in f_lower for t in terms):
                            is_match = True
                        # 3. Extension matches AND longest term matches
                        elif q_ext and f_lower.endswith(q_ext) and terms:
                            longest = max(terms, key=len)
                            if len(longest) >= 3 and longest in f_lower:
                                is_match = True

                        if is_match:
                            matches.append(os.path.join(root, file))
                        if len(matches) >= 50:
                            break
                    if len(matches) >= 50:
                        break

            # Sort by relevance score, best first
            matches.sort(key=_score, reverse=True)
            # Return top 10
            top = matches[:10]

            if top:
                return FileToolResult(True, "\n".join(top), "find_system_file")
            else:
                return FileToolResult(True, "No matching files found.", "find_system_file")
        except Exception as exc:
            return FileToolResult(False, str(exc), "find_system_file")

    def find_system_folder(self, foldername: str, search_dir: str = "~") -> FileToolResult:
        """Search for a directory by name. Prioritizes Downloads, Documents, Desktop, OneDrive, and workspace."""
        try:
            start_dir = resolve_path(search_dir) if search_dir else Path.home()
        except Exception as e:
            logger.error(f"Failed to resolve search directory '{search_dir}': {e}")
            start_dir = Path.home()
        
        try:
            if not start_dir.exists() or not start_dir.is_dir():
                return FileToolResult(False, f"Search directory invalid: {start_dir}", "find_system_folder")

            logger.info(f"🔍 Searching for folder '{foldername}' starting from: {start_dir}")

            q_lower = foldername.lower().strip()
            q_clean = q_lower.replace(" ", "")
            terms = [t for t in q_clean.split() if len(t) >= 2]

            def _score(dirpath: str) -> int:
                dname = os.path.basename(dirpath).lower()
                s = 0
                if dname == q_lower or dname == q_clean:
                    s += 1000
                if terms and all(t in dname for t in terms):
                    s += 50
                fp_lower = dirpath.lower()
                for pf in ("downloads", "documents", "desktop", "onedrive"):
                    if pf in fp_lower:
                        s += 30
                        break
                depth = dirpath.count(os.sep)
                s -= depth
                return s

            priority_names = ["Downloads", "Documents", "Desktop", "OneDrive"]
            priority_dirs = [start_dir / pn for pn in priority_names if (start_dir / pn).is_dir()]
            
            if self.workspace and self.workspace.is_dir() and self.workspace not in priority_dirs:
                priority_dirs.insert(0, self.workspace)
            
            other_dirs = [
                d for d in start_dir.iterdir()
                if d.is_dir() and d.name not in priority_names
                and d.name not in self._SKIP_DIRS and not d.name.startswith('.')
            ]
            search_order = priority_dirs + other_dirs

            matches: list[str] = []

            for search_root in search_order:
                for root, dirs, files in os.walk(str(search_root)):
                    dirs[:] = [d for d in dirs if d not in self._SKIP_DIRS and not d.startswith('.')]
                    for d in dirs:
                        d_lower = d.lower()
                        d_clean = d_lower.replace(" ", "")

                        is_match = False
                        if q_clean in d_clean or d_clean in q_clean:
                            is_match = True
                        elif terms and all(t in d_lower for t in terms):
                            is_match = True

                        if is_match:
                            matches.append(os.path.join(root, d))
                        if len(matches) >= 50:
                            break
                    if len(matches) >= 50:
                        break

            matches.sort(key=_score, reverse=True)
            top = matches[:10]

            if top:
                return FileToolResult(True, "\n".join(top), "find_system_folder")
            else:
                return FileToolResult(True, "No matching folders found.", "find_system_folder")
        except Exception as exc:
            return FileToolResult(False, str(exc), "find_system_folder")

    @staticmethod

    def _human_size(size_bytes: int) -> str:
        for unit in ["B", "KB", "MB", "GB", "TB"]:
            if size_bytes < 1024:
                return f"{size_bytes:.1f} {unit}"
            size_bytes /= 1024
        return f"{size_bytes:.1f} PB"

    def _extract_text_from_binary(self, target: Path) -> Optional[str]:
        """Attempt to extract text from known binary document formats (PDF, DOCX)."""
        ext = target.suffix.lower()
        try:
            if ext == ".pdf":
                try:
                    import pdfplumber
                    text_parts = []
                    with pdfplumber.open(str(target)) as pdf:
                        for page in pdf.pages:
                            text = page.extract_text()
                            if text:
                                text_parts.append(text)
                    return "\n\n".join(text_parts)
                except ImportError:
                    pass
                
                try:
                    from PyPDF2 import PdfReader
                    reader = PdfReader(str(target))
                    text_parts = []
                    for page in reader.pages:
                        text = page.extract_text()
                        if text:
                            text_parts.append(text)
                    return "\n\n".join(text_parts)
                except ImportError:
                    pass
            elif ext in (".docx", ".doc"):
                try:
                    from docx import Document
                    doc = Document(str(target))
                    return "\n".join(para.text for para in doc.paragraphs)
                except ImportError:
                    pass
        except Exception as e:
            logger.warning(f"Failed to extract text from {target}: {e}")
            
        return None

