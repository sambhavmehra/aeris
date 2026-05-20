"""
AERIS — File Converter Engine
Converts between file formats:
  - DOC/DOCX → PDF
  - PDF → DOCX
  - PDF → JSON (structured text extraction)
  - PDF → TXT
  - TXT → PDF
  - CSV → JSON
  - JSON → CSV
  - Markdown → HTML
  - HTML → PDF
  - Image format conversions (PNG↔JPG↔BMP↔WEBP)
  - XLSX → CSV / JSON
"""
from __future__ import annotations

import csv
import json
import logging
import os
import subprocess
import sys
from datetime import datetime
from pathlib import Path
from typing import Optional

logger = logging.getLogger("AerisFileConverter")


class FileConverter:
    """Universal file format converter for AERIS."""

    def __init__(self, output_dir: str | None = None):
        self.output_dir = Path(output_dir or os.path.join(os.getcwd(), "data", "converted"))
        self.output_dir.mkdir(parents=True, exist_ok=True)

    def convert(self, input_path: str, target_format: str, output_path: str = "") -> dict:
        """
        Main entry point — detect source format and route to the right converter.
        target_format: 'pdf', 'docx', 'json', 'txt', 'csv', 'html', 'png', 'jpg', etc.
        """
        # Sanitize: if input_path has multiple lines (e.g. from find_system_file),
        # extract just the first valid file path
        input_path = self._sanitize_path(input_path)
        if output_path:
            output_path = self._sanitize_path(output_path)

        src = Path(input_path).expanduser().resolve()
        if not src.exists():
            return {"success": False, "error": f"Source file not found: {input_path}"}

        src_ext = src.suffix.lower().lstrip(".")
        tgt = target_format.lower().strip().lstrip(".")

        if src_ext == tgt:
            return {"success": False, "error": f"Source and target formats are the same: .{tgt}"}

        # Build output path
        if not output_path:
            out_name = f"{src.stem}_converted.{tgt}"
            out = self.output_dir / out_name
        else:
            out = Path(output_path).expanduser().resolve()
            out.parent.mkdir(parents=True, exist_ok=True)

        try:
            # Route to the correct converter
            route_key = f"{src_ext}_to_{tgt}"
            converter_map = {
                # Document conversions
                "docx_to_pdf": self._docx_to_pdf,
                "doc_to_pdf": self._docx_to_pdf,
                "pdf_to_docx": self._pdf_to_docx,
                "pdf_to_doc": self._pdf_to_docx,
                "docx_to_doc": self._doc_copy,
                "doc_to_docx": self._doc_copy,
                "pdf_to_txt": self._pdf_to_txt,
                "pdf_to_json": self._pdf_to_json,
                "txt_to_pdf": self._txt_to_pdf,
                "md_to_html": self._md_to_html,
                "markdown_to_html": self._md_to_html,
                "html_to_pdf": self._html_to_pdf,
                # Spreadsheet conversions
                "csv_to_json": self._csv_to_json,
                "json_to_csv": self._json_to_csv,
                "xlsx_to_csv": self._xlsx_to_csv,
                "xlsx_to_json": self._xlsx_to_json,
                "xls_to_csv": self._xlsx_to_csv,
                # Image conversions
                "png_to_jpg": self._image_convert,
                "png_to_jpeg": self._image_convert,
                "jpg_to_png": self._image_convert,
                "jpeg_to_png": self._image_convert,
                "png_to_webp": self._image_convert,
                "webp_to_png": self._image_convert,
                "jpg_to_webp": self._image_convert,
                "webp_to_jpg": self._image_convert,
                "bmp_to_png": self._image_convert,
                "png_to_bmp": self._image_convert,
                "bmp_to_jpg": self._image_convert,
                "jpg_to_bmp": self._image_convert,
                # TXT / generic
                "txt_to_html": self._txt_to_html,
                "txt_to_json": self._txt_to_json,
                "json_to_txt": self._json_to_txt,
            }

            handler = converter_map.get(route_key)
            if handler:
                return handler(src, out)
            else:
                return {"success": False, "error": f"Unsupported conversion: .{src_ext} → .{tgt}. Supported: {', '.join(sorted(set(converter_map.keys())))}"}

        except Exception as e:
            logger.error(f"Conversion error: {e}")
            return {"success": False, "error": str(e)}

    def get_supported_conversions(self) -> list[str]:
        """Return list of supported conversion routes."""
        return [
            "docx → pdf", "doc → pdf", "pdf → docx", "pdf → txt", "pdf → json",
            "txt → pdf", "txt → html", "txt → json", "md → html", "html → pdf",
            "csv → json", "json → csv", "xlsx → csv", "xlsx → json",
            "png ↔ jpg", "png ↔ webp", "jpg ↔ webp", "bmp ↔ png", "bmp ↔ jpg",
            "json → txt",
        ]

    # ═════════════════════════════════════════════════════════════════
    #  DOCUMENT CONVERTERS
    # ═════════════════════════════════════════════════════════════════

    def _doc_copy(self, src: Path, out: Path) -> dict:
        """Handle copy/rename between DOC and DOCX formats."""
        import shutil
        try:
            shutil.copy2(src, out)
            return {"success": True, "output_path": str(out), "method": "copy", "size_bytes": out.stat().st_size}
        except Exception as e:
            return {"success": False, "error": f"Document copy failed: {e}"}

    def _docx_to_pdf(self, src: Path, out: Path) -> dict:
        """Convert DOCX/DOC to PDF using python-docx + reportlab, or LibreOffice fallback."""
        # Try LibreOffice first (best quality)
        try:
            out_dir = str(out.parent)
            result = subprocess.run(
                ["soffice", "--headless", "--convert-to", "pdf", "--outdir", out_dir, str(src)],
                capture_output=True, text=True, timeout=120,
            )
            # LibreOffice names the output based on input filename
            lo_output = out.parent / f"{src.stem}.pdf"
            if lo_output.exists():
                if lo_output != out:
                    lo_output.rename(out)
                return {"success": True, "output_path": str(out), "method": "libreoffice", "size_bytes": out.stat().st_size}
        except (FileNotFoundError, subprocess.TimeoutExpired):
            pass

        # Fallback: python-docx + reportlab
        try:
            from docx import Document
            from reportlab.lib.pagesizes import letter
            from reportlab.lib.styles import getSampleStyleSheet
            from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
            from reportlab.lib.units import inch

            doc = Document(str(src))
            pdf = SimpleDocTemplate(str(out), pagesize=letter)
            styles = getSampleStyleSheet()
            story = []

            for para in doc.paragraphs:
                text = para.text.strip()
                if not text:
                    story.append(Spacer(1, 0.2 * inch))
                    continue
                # Detect heading styles
                if para.style.name.startswith("Heading"):
                    story.append(Paragraph(text, styles["Heading1"]))
                else:
                    story.append(Paragraph(text, styles["Normal"]))
                story.append(Spacer(1, 0.05 * inch))

            pdf.build(story)
            return {"success": True, "output_path": str(out), "method": "python-docx+reportlab", "size_bytes": out.stat().st_size}
        except ImportError as ie:
            return {"success": False, "error": f"Missing dependency: {ie}. Install with: pip install python-docx reportlab"}
        except Exception as e:
            return {"success": False, "error": f"DOCX→PDF failed: {e}"}

    def _pdf_to_docx(self, src: Path, out: Path) -> dict:
        """Convert PDF to DOCX using pdf2docx."""
        out = out.with_suffix(".docx")
        try:
            from pdf2docx import Converter
            cv = Converter(str(src))
            cv.convert(str(out), start=0, end=None)
            cv.close()
            return {"success": True, "output_path": str(out), "method": "pdf2docx", "size_bytes": out.stat().st_size}
        except ImportError:
            # Fallback: extract text and create a basic docx
            try:
                text = self._extract_pdf_text(src)
                from docx import Document
                doc = Document()
                for line in text.split("\n"):
                    if line.strip():
                        doc.add_paragraph(line)
                doc.save(str(out))
                return {"success": True, "output_path": str(out), "method": "pdfplumber+python-docx", "size_bytes": out.stat().st_size}
            except ImportError as ie:
                return {"success": False, "error": f"Missing dependency: {ie}. Install with: pip install pdf2docx python-docx pdfplumber"}
        except Exception as e:
            return {"success": False, "error": f"PDF→DOCX failed: {e}"}

    def _pdf_to_txt(self, src: Path, out: Path) -> dict:
        """Extract text from PDF."""
        out = out.with_suffix(".txt")
        try:
            text = self._extract_pdf_text(src)
            out.write_text(text, encoding="utf-8")
            return {"success": True, "output_path": str(out), "method": "pdfplumber", "size_bytes": out.stat().st_size, "char_count": len(text)}
        except ImportError as ie:
            return {"success": False, "error": f"Missing dependency: {ie}. Install with: pip install pdfplumber"}
        except Exception as e:
            return {"success": False, "error": f"PDF→TXT failed: {e}"}

    def _pdf_to_json(self, src: Path, out: Path) -> dict:
        """Extract structured text from PDF into JSON format."""
        out = out.with_suffix(".json")
        try:
            import pdfplumber
            data = {"source_file": str(src), "pages": [], "total_pages": 0, "extracted_at": datetime.now().isoformat()}

            with pdfplumber.open(str(src)) as pdf:
                data["total_pages"] = len(pdf.pages)
                for i, page in enumerate(pdf.pages):
                    page_data = {
                        "page_number": i + 1,
                        "width": page.width,
                        "height": page.height,
                        "text": page.extract_text() or "",
                    }
                    # Try extracting tables
                    tables = page.extract_tables()
                    if tables:
                        page_data["tables"] = []
                        for table in tables:
                            page_data["tables"].append(table)
                    data["pages"].append(page_data)

            out.write_text(json.dumps(data, indent=2, ensure_ascii=False), encoding="utf-8")
            return {"success": True, "output_path": str(out), "method": "pdfplumber", "pages": data["total_pages"], "size_bytes": out.stat().st_size}
        except ImportError:
            return {"success": False, "error": "Missing dependency: pdfplumber. Install with: pip install pdfplumber"}
        except Exception as e:
            return {"success": False, "error": f"PDF→JSON failed: {e}"}

    def _txt_to_pdf(self, src: Path, out: Path) -> dict:
        """Convert plain text to PDF."""
        out = out.with_suffix(".pdf")
        try:
            from reportlab.lib.pagesizes import letter
            from reportlab.lib.styles import getSampleStyleSheet
            from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
            from reportlab.lib.units import inch

            text = src.read_text(encoding="utf-8", errors="replace")
            pdf = SimpleDocTemplate(str(out), pagesize=letter)
            styles = getSampleStyleSheet()
            story = []

            for line in text.split("\n"):
                if line.strip():
                    # Escape XML special chars for reportlab
                    safe_line = line.replace("&", "&amp;").replace("<", "&lt;").replace(">", "&gt;")
                    story.append(Paragraph(safe_line, styles["Normal"]))
                else:
                    story.append(Spacer(1, 0.15 * inch))

            if not story:
                story.append(Paragraph("(empty document)", styles["Normal"]))

            pdf.build(story)
            return {"success": True, "output_path": str(out), "method": "reportlab", "size_bytes": out.stat().st_size}
        except ImportError as ie:
            return {"success": False, "error": f"Missing dependency: {ie}. Install with: pip install reportlab"}
        except Exception as e:
            return {"success": False, "error": f"TXT→PDF failed: {e}"}

    def _md_to_html(self, src: Path, out: Path) -> dict:
        """Convert Markdown to styled HTML."""
        out = out.with_suffix(".html")
        try:
            import markdown
            md_text = src.read_text(encoding="utf-8", errors="replace")
            html_body = markdown.markdown(md_text, extensions=["tables", "fenced_code", "codehilite"])
        except ImportError:
            # Fallback: basic conversion
            md_text = src.read_text(encoding="utf-8", errors="replace")
            html_body = self._basic_md_to_html(md_text)

        html = f"""<!DOCTYPE html>
<html lang="en">
<head>
  <meta charset="UTF-8"/>
  <meta name="viewport" content="width=device-width, initial-scale=1.0"/>
  <title>{src.stem}</title>
  <style>
    body {{ font-family: 'Inter', system-ui, sans-serif; max-width: 800px; margin: 40px auto; padding: 0 24px; background: #0a0a0f; color: #e8e6f0; line-height: 1.7; }}
    h1, h2, h3 {{ background: linear-gradient(135deg, #8b5cf6, #06b6d4); -webkit-background-clip: text; -webkit-text-fill-color: transparent; }}
    code {{ background: rgba(139,92,246,.1); padding: 2px 6px; border-radius: 4px; font-size: 0.9em; }}
    pre {{ background: rgba(15,15,25,.9); border: 1px solid rgba(139,92,246,.15); border-radius: 8px; padding: 16px; overflow-x: auto; }}
    table {{ border-collapse: collapse; width: 100%; margin: 16px 0; }}
    th, td {{ border: 1px solid rgba(139,92,246,.2); padding: 8px 12px; text-align: left; }}
    th {{ background: rgba(139,92,246,.1); }}
    a {{ color: #8b5cf6; }}
    blockquote {{ border-left: 3px solid #8b5cf6; margin: 16px 0; padding: 8px 16px; color: #9896a8; }}
  </style>
</head>
<body>
{html_body}
</body>
</html>"""
        out.write_text(html, encoding="utf-8")
        return {"success": True, "output_path": str(out), "method": "markdown", "size_bytes": out.stat().st_size}

    def _html_to_pdf(self, src: Path, out: Path) -> dict:
        """Convert HTML to PDF."""
        out = out.with_suffix(".pdf")
        
        # Try playwright (since it's highly robust on Windows/any platform and works out-of-the-box in this env)
        try:
            from playwright.sync_api import sync_playwright
            with sync_playwright() as p:
                browser = p.chromium.launch(headless=True)
                page = browser.new_page()
                # Use file:/// URL for the source file to resolve local resources/images if any
                src_url = src.resolve().as_uri()
                page.goto(src_url)
                # Print to PDF with margin and formatting
                page.pdf(path=str(out), format="Letter", print_background=True)
                browser.close()
            return {"success": True, "output_path": str(out), "method": "playwright", "size_bytes": out.stat().st_size}
        except Exception as e:
            logger.warning(f"Playwright HTML→PDF failed: {e}. Trying fallbacks.")
            pass

        # Try weasyprint
        try:
            from weasyprint import HTML
            HTML(filename=str(src)).write_pdf(str(out))
            return {"success": True, "output_path": str(out), "method": "weasyprint", "size_bytes": out.stat().st_size}
        except ImportError:
            pass

        # Try pdfkit
        try:
            import pdfkit
            pdfkit.from_file(str(src), str(out))
            return {"success": True, "output_path": str(out), "method": "pdfkit", "size_bytes": out.stat().st_size}
        except ImportError:
            pass

        return {"success": False, "error": "Missing dependency for HTML→PDF. Install with: pip install playwright OR weasyprint OR pdfkit"}

    # ═════════════════════════════════════════════════════════════════
    #  SPREADSHEET / DATA CONVERTERS
    # ═════════════════════════════════════════════════════════════════

    def _csv_to_json(self, src: Path, out: Path) -> dict:
        """Convert CSV to JSON array."""
        out = out.with_suffix(".json")
        try:
            with open(src, newline="", encoding="utf-8-sig") as f:
                reader = csv.DictReader(f)
                rows = list(reader)
            out.write_text(json.dumps(rows, indent=2, ensure_ascii=False), encoding="utf-8")
            return {"success": True, "output_path": str(out), "method": "csv.DictReader", "row_count": len(rows), "size_bytes": out.stat().st_size}
        except Exception as e:
            return {"success": False, "error": f"CSV→JSON failed: {e}"}

    def _json_to_csv(self, src: Path, out: Path) -> dict:
        """Convert JSON array to CSV."""
        out = out.with_suffix(".csv")
        try:
            data = json.loads(src.read_text(encoding="utf-8"))
            if not isinstance(data, list) or not data:
                return {"success": False, "error": "JSON must be a non-empty array of objects"}

            fieldnames = list(data[0].keys())
            with open(out, "w", newline="", encoding="utf-8") as f:
                writer = csv.DictWriter(f, fieldnames=fieldnames)
                writer.writeheader()
                for row in data:
                    writer.writerow(row)

            return {"success": True, "output_path": str(out), "method": "csv.DictWriter", "row_count": len(data), "size_bytes": out.stat().st_size}
        except Exception as e:
            return {"success": False, "error": f"JSON→CSV failed: {e}"}

    def _xlsx_to_csv(self, src: Path, out: Path) -> dict:
        """Convert Excel XLSX to CSV."""
        out = out.with_suffix(".csv")
        try:
            import openpyxl
            wb = openpyxl.load_workbook(str(src), read_only=True, data_only=True)
            ws = wb.active
            with open(out, "w", newline="", encoding="utf-8") as f:
                writer = csv.writer(f)
                for row in ws.iter_rows(values_only=True):
                    writer.writerow(row)
            wb.close()
            return {"success": True, "output_path": str(out), "method": "openpyxl", "size_bytes": out.stat().st_size}
        except ImportError:
            return {"success": False, "error": "Missing dependency: openpyxl. Install with: pip install openpyxl"}
        except Exception as e:
            return {"success": False, "error": f"XLSX→CSV failed: {e}"}

    def _xlsx_to_json(self, src: Path, out: Path) -> dict:
        """Convert Excel XLSX to JSON."""
        out = out.with_suffix(".json")
        try:
            import openpyxl
            wb = openpyxl.load_workbook(str(src), read_only=True, data_only=True)
            ws = wb.active
            rows = list(ws.iter_rows(values_only=True))
            wb.close()

            if not rows:
                return {"success": False, "error": "Empty spreadsheet"}

            headers = [str(h) if h else f"col_{i}" for i, h in enumerate(rows[0])]
            data = []
            for row in rows[1:]:
                row_dict = {}
                for i, val in enumerate(row):
                    key = headers[i] if i < len(headers) else f"col_{i}"
                    row_dict[key] = val
                data.append(row_dict)

            out.write_text(json.dumps(data, indent=2, default=str, ensure_ascii=False), encoding="utf-8")
            return {"success": True, "output_path": str(out), "method": "openpyxl", "row_count": len(data), "size_bytes": out.stat().st_size}
        except ImportError:
            return {"success": False, "error": "Missing dependency: openpyxl. Install with: pip install openpyxl"}
        except Exception as e:
            return {"success": False, "error": f"XLSX→JSON failed: {e}"}

    # ═════════════════════════════════════════════════════════════════
    #  IMAGE CONVERTERS
    # ═════════════════════════════════════════════════════════════════

    def _image_convert(self, src: Path, out: Path) -> dict:
        """Convert between image formats using Pillow."""
        try:
            from PIL import Image
            img = Image.open(str(src))

            # Handle transparency for JPEG
            tgt_ext = out.suffix.lower()
            if tgt_ext in (".jpg", ".jpeg", ".bmp") and img.mode in ("RGBA", "LA", "P"):
                background = Image.new("RGB", img.size, (255, 255, 255))
                if img.mode == "P":
                    img = img.convert("RGBA")
                background.paste(img, mask=img.split()[-1] if "A" in img.mode else None)
                img = background

            img.save(str(out))
            return {
                "success": True,
                "output_path": str(out),
                "method": "pillow",
                "original_size": f"{img.size[0]}x{img.size[1]}",
                "size_bytes": out.stat().st_size,
            }
        except ImportError:
            return {"success": False, "error": "Missing dependency: Pillow. Install with: pip install Pillow"}
        except Exception as e:
            return {"success": False, "error": f"Image conversion failed: {e}"}

    # ═════════════════════════════════════════════════════════════════
    #  TEXT / GENERIC CONVERTERS
    # ═════════════════════════════════════════════════════════════════

    def _txt_to_html(self, src: Path, out: Path) -> dict:
        """Convert plain text to styled HTML."""
        out = out.with_suffix(".html")
        text = src.read_text(encoding="utf-8", errors="replace")
        paragraphs = "\n".join(
            f"<p>{line}</p>" if line.strip() else "<br/>"
            for line in text.split("\n")
        )
        html = f"""<!DOCTYPE html>
<html lang="en">
<head>
  <meta charset="UTF-8"/>
  <title>{src.stem}</title>
  <style>
    body {{ font-family: 'Inter', system-ui, sans-serif; max-width: 800px; margin: 40px auto; padding: 0 24px; background: #0a0a0f; color: #e8e6f0; line-height: 1.7; }}
  </style>
</head>
<body>
{paragraphs}
</body>
</html>"""
        out.write_text(html, encoding="utf-8")
        return {"success": True, "output_path": str(out), "method": "text-to-html", "size_bytes": out.stat().st_size}

    def _txt_to_json(self, src: Path, out: Path) -> dict:
        """Convert text file to JSON with line-by-line structure."""
        out = out.with_suffix(".json")
        text = src.read_text(encoding="utf-8", errors="replace")
        lines = text.split("\n")
        data = {
            "source": str(src),
            "line_count": len(lines),
            "content": text,
            "lines": [{"line_number": i + 1, "text": line} for i, line in enumerate(lines)],
        }
        out.write_text(json.dumps(data, indent=2, ensure_ascii=False), encoding="utf-8")
        return {"success": True, "output_path": str(out), "method": "text-to-json", "line_count": len(lines), "size_bytes": out.stat().st_size}

    def _json_to_txt(self, src: Path, out: Path) -> dict:
        """Convert JSON to pretty-printed text."""
        out = out.with_suffix(".txt")
        data = json.loads(src.read_text(encoding="utf-8"))
        text = json.dumps(data, indent=2, ensure_ascii=False, default=str)
        out.write_text(text, encoding="utf-8")
        return {"success": True, "output_path": str(out), "method": "json-to-txt", "size_bytes": out.stat().st_size}

    # ═════════════════════════════════════════════════════════════════
    #  HELPERS
    # ═════════════════════════════════════════════════════════════════

    @staticmethod
    def _sanitize_path(raw: str) -> str:
        """Extract the first valid file path from a string.
        Handles multi-line input from find_system_file or $prev injection."""
        import re
        raw = raw.strip()
        # If no newlines, it's already a single path
        if "\n" not in raw:
            return raw
        # Pick the first line that looks like a real path
        for line in raw.splitlines():
            line = line.strip()
            if not line:
                continue
            # Windows absolute path
            if re.match(r'^[A-Za-z]:[\\\/]', line):
                return line
            # Unix absolute path
            if line.startswith("/") or line.startswith("~"):
                return line
        # Fallback: first non-empty line
        for line in raw.splitlines():
            if line.strip():
                return line.strip()
        return raw

    def _extract_pdf_text(self, src: Path) -> str:
        """Extract text from a PDF using pdfplumber (primary) or PyPDF2 (fallback)."""
        try:
            import pdfplumber
            text_parts = []
            with pdfplumber.open(str(src)) as pdf:
                for page in pdf.pages:
                    page_text = page.extract_text()
                    if page_text:
                        text_parts.append(page_text)
            return "\n\n".join(text_parts)
        except ImportError:
            pass

        # Fallback: PyPDF2
        try:
            from PyPDF2 import PdfReader
            reader = PdfReader(str(src))
            text_parts = []
            for page in reader.pages:
                text = page.extract_text()
                if text:
                    text_parts.append(text)
            return "\n\n".join(text_parts)
        except ImportError:
            raise ImportError("Install pdfplumber or PyPDF2: pip install pdfplumber PyPDF2")

    def _basic_md_to_html(self, md_text: str) -> str:
        """Basic markdown → HTML without external library."""
        import re
        lines = md_text.split("\n")
        html_lines = []
        in_code_block = False

        for line in lines:
            if line.startswith("```"):
                if in_code_block:
                    html_lines.append("</code></pre>")
                    in_code_block = False
                else:
                    html_lines.append("<pre><code>")
                    in_code_block = True
                continue

            if in_code_block:
                html_lines.append(line)
                continue

            # Headers
            if line.startswith("### "):
                html_lines.append(f"<h3>{line[4:]}</h3>")
            elif line.startswith("## "):
                html_lines.append(f"<h2>{line[3:]}</h2>")
            elif line.startswith("# "):
                html_lines.append(f"<h1>{line[2:]}</h1>")
            elif line.startswith("- ") or line.startswith("* "):
                html_lines.append(f"<li>{line[2:]}</li>")
            elif line.startswith("> "):
                html_lines.append(f"<blockquote>{line[2:]}</blockquote>")
            elif line.strip() == "":
                html_lines.append("<br/>")
            else:
                # Inline code
                line = re.sub(r"`([^`]+)`", r"<code>\1</code>", line)
                # Bold
                line = re.sub(r"\*\*([^*]+)\*\*", r"<strong>\1</strong>", line)
                # Italic
                line = re.sub(r"\*([^*]+)\*", r"<em>\1</em>", line)
                html_lines.append(f"<p>{line}</p>")

        if in_code_block:
            html_lines.append("</code></pre>")

        return "\n".join(html_lines)
